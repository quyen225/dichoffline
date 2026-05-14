# main.py
import os
import sys
import re
import json
import time
import threading
import customtkinter as ctk
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Thread
from tkinter import filedialog, messagebox

# --- GIẢI PHÁP ĐĂNG KÝ PATH TỰ ĐỘNG ---
# Lấy đường dẫn tuyệt đối của thư mục gốc và thư mục model
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODEL_DIR = os.path.join(BASE_DIR, "model")

# Ép Python phải tìm kiếm trong cả thư mục gốc và thư mục model khi thực hiện lệnh import
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)
if MODEL_DIR not in sys.path:
    sys.path.insert(0, MODEL_DIR)

# Bây giờ bạn có thể import theo bất kỳ cách nào mà không sợ bị lỗi No module named
try:
    from ollama_provider import OllamaProvider
    from groq_provider import GroqProvider
    from openai_provider import OpenAIProvider
    from azure_provider import AzureProvider
except ImportError:
    try:
        from model.ollama_provider import OllamaProvider
        from model.groq_provider import GroqProvider
        from model.openai_provider import OpenAIProvider
        from model.azure_provider import AzureProvider
    except ImportError as e:
        print(f"Lỗi Import nhà cung cấp AI: {e}")
        input("Nhấn Enter để thoát...")
        sys.exit()

try:
    from gui_manager import ApiManagerWindow
    import utils
    from utils import call_with_retry
except ImportError as e:
    print(f"Lỗi Import module hệ thống: {e}")
    input("Kiểm tra lại các file .py và nhấn Enter để thoát...")
    sys.exit()

class TranslatorApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("AI Translator Master - Professional Hybrid v3.5 (Worker Pools)")
        self.geometry("1400x1050")

        # --- State quản lý ---
        self.api_window = None
        self.is_running = False
        self.content_lines = []
        self.translated_lines = []
        self.name_dict = {}
        self.current_filename = ""

        self.setup_ui()
        self.refresh_ollama()

    def log(self, msg):
        ts = time.strftime("%Y-%m-%d %H:%M:%S")
        try:
            self.log_box.insert("end", f"[{ts}] {msg}\n")
            self.log_box.see("end")
        except Exception:
            print(f"[{ts}] {msg}")

    def setup_ui(self):
        # 1. Toolbar trên cùng
        top = ctk.CTkFrame(self)
        top.pack(fill="x", padx=20, pady=10)
        ctk.CTkButton(top, text="🔧 Cấu hình API", command=self.open_api, width=150).pack(side="left", padx=10)

        ctk.CTkLabel(top, text="Model Ollama:").pack(side="left", padx=5)
        self.model_selector = ctk.CTkOptionMenu(top, values=["Đang quét..."], width=200)
        self.model_selector.pack(side="left", padx=5)

        self.status_engine_label = ctk.CTkLabel(top, text="⚙️ Hệ thống: Sẵn sàng", font=("Arial", 13, "bold"), text_color="#ABB2B9")
        self.status_engine_label.pack(side="left", padx=20)

        ctk.CTkButton(top, text="📁 Nạp Truyện", fg_color="#2980B9", command=self.load_story).pack(side="right", padx=10)

        # 2. Hàng Từ điển
        dict_frame = ctk.CTkFrame(self)
        dict_frame.pack(fill="x", padx=20, pady=5)
        self.btn_dict = ctk.CTkButton(dict_frame, text="📁 Nạp Từ Điển (.txt)", width=200, fg_color="#16A085", command=self.load_dict)
        self.btn_dict.pack(side="left", padx=10)
        self.dict_status_label = ctk.CTkLabel(dict_frame, text="Chưa nạp từ điển", text_color="gray")
        self.dict_status_label.pack(side="left", padx=10)

        # 3. KHU VỰC DỊCH THỬ
        test_frame = ctk.CTkFrame(self)
        test_frame.pack(fill="x", padx=20, pady=5)
        ctk.CTkLabel(test_frame, text="Dịch thử:").pack(side="left", padx=10)
        self.test_input = ctk.CTkEntry(test_frame, placeholder_text="Nhập đoạn tiếng Trung ngắn để test...", width=600)
        self.test_input.pack(side="left", padx=5, fill="x", expand=True)
        ctk.CTkButton(test_frame, text="Test Ollama", width=100, fg_color="#3498DB", command=lambda: self.test_quick_api("ollama")).pack(side="left", padx=5)
        ctk.CTkButton(test_frame, text="Test Groq", width=100, fg_color="#E67E22", command=lambda: self.test_quick_api("groq")).pack(side="left", padx=5)
        self.test_res_label = ctk.CTkLabel(self, text="", font=("Arial", 12, "italic"), text_color="#F1C40F")
        self.test_res_label.pack(pady=2)

        # 4. Log Box
        self.log_box = ctk.CTkTextbox(self, height=180, fg_color="black", text_color="#00FF00", font=("Consolas", 12))
        self.log_box.pack(fill="x", padx=20, pady=10)

        # 5. Tiến độ
        prog_frame = ctk.CTkFrame(self, fg_color="transparent")
        prog_frame.pack(fill="x", padx=25)
        self.progress_label = ctk.CTkLabel(prog_frame, text="Tiến độ: 0% (Dòng 0/0)", font=("Arial", 12))
        self.progress_label.pack(side="left")
        self.progress_bar = ctk.CTkProgressBar(self)
        self.progress_bar.set(0)
        self.progress_bar.pack(fill="x", padx=20, pady=5)

        # 6. Hàng nút điều khiển chính
        ctrl_frame = ctk.CTkFrame(self, fg_color="transparent")
        ctrl_frame.pack(pady=15)
        ctk.CTkButton(ctrl_frame, text="🚀 DỊCH MỚI", width=160, height=45, command=lambda: self.start_task("new")).pack(side="left", padx=8)
        ctk.CTkButton(ctrl_frame, text="🩹 CỨU HỘ", width=160, height=45, fg_color="#D35400", command=lambda: self.start_task("rescue")).pack(side="left", padx=8)
        ctk.CTkButton(ctrl_frame, text="🛑 DỪNG", width=100, height=45, fg_color="#C0392B", command=self.stop_task).pack(side="left", padx=8)
        ctk.CTkButton(ctrl_frame, text="💾 XUẤT FILE", width=160, height=45, fg_color="#27AE60", command=self.handle_export).pack(side="left", padx=8)

        # 7. Output Preview
        self.output_view = ctk.CTkTextbox(self, height=350, fg_color="#1C1C1C", font=("Arial", 14))
        self.output_view.pack(fill="both", expand=True, padx=20, pady=10)

    # ================= LOGIC ĐIỀU PHỐI VÀ XỬ LÝ ĐA LUỒNG =================

    def load_dict(self):
        path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
        if path:
            self.name_dict = {}
            try:
                with open(path, "r", encoding="utf-8") as f:
                    for line in f:
                        if "=" in line:
                            k, v = line.strip().split("=", 1)
                            self.name_dict[k.strip()] = v.strip()
                self.dict_status_label.configure(text=f"✅ Đã nạp: {os.path.basename(path)} ({len(self.name_dict)} từ)", text_color="#2ECC71")
                self.log(f"📚 Đã nạp từ điển: {len(self.name_dict)} từ.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Lỗi nạp từ điển: {e}")

    def load_story(self):
        path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
        if not path: 
            return
        self.current_filename = os.path.basename(path)
        try:
            with open(path, "r", encoding="utf-8") as f:
                self.content_lines = [line.strip() for line in f if line.strip()]
            self.translated_lines = [None] * len(self.content_lines)

            old_data = utils.load_autosave(self.current_filename)
            if old_data and len(old_data) == len(self.content_lines):
                if messagebox.askyesno("Khôi phục", "Tìm thấy bản dịch tạm. Khôi phục lại?"):
                    self.translated_lines = old_data
                    self.log("♻️ Đã khôi phục dữ liệu từ bản lưu tự động.")
            
            self.log(f"✅ Đã nạp {len(self.content_lines)} dòng truyện.")
            self.update_ui_initial()
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

    def test_quick_api(self, engine_type):
        text = self.test_input.get().strip()
        if not text: 
            return
        def task():
            self.log(f"🧪 Test bằng {engine_type.upper()}...")
            try:
                processed = utils.replace_names(text, self.name_dict)
                if engine_type == "ollama":
                    model = self.model_selector.get()
                    res = call_with_retry(lambda: OllamaProvider.generate(model, processed), retries=2, initial_backoff=0.8)
                else:
                    g_key = os.getenv("GROQ_KEY")
                    g_model = self.api_window.groq_model_selector.get() if (self.api_window and self.api_window.winfo_exists()) else "llama-3.3-70b-versatile"
                    res = call_with_retry(lambda: GroqProvider.generate(g_key, g_model, processed), retries=2, initial_backoff=1.0)
                color = "#2ECC71" if not (isinstance(res, str) and res.startswith("[Lỗi")) else "red"
                self.after(0, lambda: self.test_res_label.configure(text=f"Kết quả: {res}", text_color=color))
            except Exception as e:
                self.log(f"❌ Lỗi: {e}")
        Thread(target=task, daemon=True).start()

    def start_task(self, mode):
        if not self.content_lines:
            return messagebox.showwarning("Lỗi", "Chưa nạp truyện!")
        if self.is_running:
            return
        if mode == "new":
            self.output_view.delete("1.0", "end")
        self.is_running = True
        # Điều phối chính trên một luồng riêng biệt để tránh đơ giao diện UI
        Thread(target=self.translation_pool_coordinator, args=(mode,), daemon=True).start()

    def translate_single_line_worker(self, i, src, mode):
        """Worker độc lập dịch thuật cho từng dòng truyện cụ thể"""
        if not self.is_running:
            return i, None, "Dừng bởi người dùng"
        try:
            # Chế độ cứu hộ sử dụng Groq trực tuyến
            if mode == "rescue":
                g_key = os.getenv("GROQ_KEY")
                g_model = self.api_window.groq_model_selector.get() if (self.api_window and self.api_window.winfo_exists()) else "llama-3.3-70b-versatile"
                result = call_with_retry(lambda: GroqProvider.generate(g_key, g_model, src),
                                         retries=4, initial_backoff=1.0, max_backoff=8.0, exceptions=(Exception,))
            # Chế độ dịch mới ưu tiên Ollama local offline
            else:
                model = self.model_selector.get()
                try:
                    result = call_with_retry(lambda: OllamaProvider.generate(model, src),
                                             retries=3, initial_backoff=0.8, max_backoff=6.0, exceptions=(Exception,))
                except Exception as e_oll:
                    self.log(f"⚠️ Ollama lỗi: {e_oll}")
                    g_key = os.getenv("GROQ_KEY")
                    if g_key:
                        try:
                            g_model = self.api_window.groq_model_selector.get() if (self.api_window and self.api_window.winfo_exists()) else "llama-3.3-70b-versatile"
                            result = call_with_retry(lambda: GroqProvider.generate(g_key, g_model, src),
                                                     retries=4, initial_backoff=1.0, max_backoff=8.0, exceptions=(Exception,))
                        except Exception as e_groq:
                            raise RuntimeError(f"Groq fallback cũng lỗi: {e_groq}") from e_groq
                    else:
                        raise

            if result:
                result = re.sub(r'\(Note:.*?\)|Note:.*', '', result, flags=re.DOTALL | re.IGNORECASE).strip()
                result = re.sub(r'^(Bản dịch|Vietnamese|Translation|Dịch):\s*', '', result, flags=re.IGNORECASE).strip()

            if hasattr(utils, "postprocess_translation"):
                result = utils.postprocess_translation(result, self.name_dict)

            return i, result, None
        except Exception as e:
            return i, f"[Lỗi dịch] Dòng {i+1}: {e}", str(e)

    def translation_pool_coordinator(self, mode):
        """Bộ điều phối trung tâm quản lý Worker Pools đa luồng"""
        total = len(self.content_lines)
        if mode == "new":
            indices = list(range(total))
        else:
            indices = [i for i, r in enumerate(self.translated_lines) if r is None or utils.count_hanzi(r) > 0]

        self.log(f"⚡ Khởi chạy Worker Pools ({'Mới' if mode=='new' else 'Cứu hộ'}) - Tổng {len(indices)} mục.")
        
                # Cấu hình luồng: 2 cho local Ollama, 3 cho API Groq
        max_workers = 3 if mode == "rescue" else 2
        autosave_every = 10
        autosave_time_interval = 30
        last_autosave_time = time.time()
        
        counter_lock = threading.Lock()
        idx_count = 0

        if mode == "rescue":
            self.after(0, lambda: self.status_engine_label.configure(text="🌐 Hệ thống: Groq Online Pools", text_color="#E67E22"))
        else:
            self.after(0, lambda: self.status_engine_label.configure(text="⚙️ Hệ thống: Ollama Offline Pools", text_color="#3498DB"))

        # Lấy thời gian cooldown từ đúng biến giao diện gốc của ứng dụng
        try:
            delay_val = float(self.cooldown_slider.get())
        except (AttributeError, NameError):
            delay_val = 0.2

        # Khởi chạy cụm Worker Pool song song quản lý luồng
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(
                    self.translate_single_line_worker, 
                    i, 
                    utils.replace_names(self.content_lines[i], self.name_dict), 
                    mode
                ): i for i in indices
            }

            for future in as_completed(futures):
                if not self.is_running:
                    executor.shutdown(wait=False, cancel_futures=True)
                    self.log("⏸️ Tiến trình dịch đa luồng bị dừng bởi người dùng.")
                    break

                i, result, error = future.result()
                
                with counter_lock:
                    idx_count += 1

                if error:
                    self.translated_lines[i] = result
                    self.log(result)
                else:
                    self.translated_lines[i] = result or ""
                    self.log(f"✅ Dòng {i+1}/{total} hoàn tất (Đa luồng).")

                # Cập nhật UI an toàn trên luồng chính (Main Thread)
                self.after(0, lambda idx=i, tot=total, txt=self.translated_lines[i]: self.update_ui_append_line(idx, tot, txt))

                # Cơ chế tự động sao lưu dữ liệu tạm thời
                if idx_count % autosave_every == 0 or (time.time() - last_autosave_time) > autosave_time_interval:
                    try:
                        utils.save_autosave(self.current_filename, self.translated_lines)
                        self.log("💾 Đã tự động lưu bản sao tạm (autosave đa luồng).")
                    except Exception as e:
                        self.log(f"⚠️ Lỗi autosave đa luồng: {e}")
                    last_autosave_time = time.time()

                # Áp dụng độ trễ nghỉ luồng động dựa trên cấu hình slider
                if delay_val > 0:
                    time.sleep(delay_val)

        self.is_running = False
        try:
            utils.save_autosave(self.current_filename, self.translated_lines)
            self.log("💾 Bản lưu đa luồng cuối cùng hoàn tất.")
        except Exception as e:
            self.log(f"⚠️ Lỗi lưu cuối: {e}")

        self.after(0, lambda: self.status_engine_label.configure(text="✅ Hoàn thành", text_color="#2ECC71"))
        self.log("🏁 Xử lý Worker Pools hoàn tất.")


    def update_ui_initial(self):
        total = len(self.content_lines)
        self.progress_bar.set(0)
        self.progress_label.configure(text=f"Tiến độ: 0% (Dòng 0/{total})")
        self.status_engine_label.configure(text="⚙️ Hệ thống: Sẵn sàng")
        self.output_view.delete("1.0", "end")
        for ln in self.translated_lines:
            if ln:
                self.output_view.insert("end", ln + "\n\n")
        self.output_view.see("end")

    def update_ui_append_line(self, index: int, total: int, line_text: str):
        pct = int(((index + 1) / total) * 100) if total else 100
        self.progress_bar.set((index + 1) / total if total else 1.0)
        self.progress_label.configure(text=f"Tiến độ: {pct}% (Dòng {index+1}/{total})")
        if line_text is None:
            line_text = ""
        self.output_view.insert("end", line_text + "\n\n")
        self.output_view.see("end")

    def render_full_output(self):
        self.output_view.delete("1.0", "end")
        for line in self.translated_lines:
            self.output_view.insert("end", (line or "") + "\n\n")
        self.output_view.see("end")

    # ================= VALIDATE API KEYS =================

    def validate_ai(self, name, env_name, entry_widget):
        key = entry_widget.get().strip()
        if not key:
            self.log(f"⚠️ Vui lòng nhập Key cho {name}")
            return

        # Nạp trực tiếp hàm đọc cấu trúc file từ thư mục model/config.py
        from model.config import get_models_by_provider

        def task():
            self.after(0, lambda: self.log(f"🔍 Đang xác thực cấu hình {name} nội bộ..."))
            try:
                # Ép hệ thống đọc danh sách model từ file cục bộ thay vì gửi request qua API mạng
                if "Groq" in name:
                    models = get_models_by_provider("groq")
                elif "OpenAI" in name:
                    models = get_models_by_provider("openai")
                else:
                    models = ["ok"]

                # Cấu hình dự phòng nếu file config cục bộ bị trống
                if not models:
                    models = ["llama-3.3-70b-versatile"]

                # Lưu khóa cấu hình vào file bảo mật .env và nạp vào bộ nhớ hệ thống
                from dotenv import set_key
                set_key(".env", env_name, key)
                os.environ[env_name] = key

                def success_ui():
                    # Đổ trực tiếp danh sách đọc từ file nội bộ vào dropdown trên UI quản lý
                    if "Groq" in name and self.api_window:
                        self.api_window.update_groq_dropdown(models)
                    self.log(f"✅ {name} OK (Đã nạp cấu trúc file nội bộ)!")
                    messagebox.showinfo("OK", f"Cấu hình {name} đã được lưu và kích hoạt thành công!")
                
                self.after(0, success_ui)
                
            except Exception as e:
                error_msg = str(e)
                self.after(0, lambda: self.log(f"❌ {name} lỗi cấu hình: {error_msg}"))
                
        Thread(target=task, daemon=True).start()


    def validate_azure(self, key_ent, ep_ent, reg_combo):
        key, ep, reg = key_ent.get().strip(), ep_ent.get().strip(), reg_combo.get()
        if not key or not ep:
            self.log("⚠️ Vui lòng nhập đầy đủ Azure Key và Endpoint.")
            return
        def task():
            try:
                res = AzureProvider.generate(key, ep, reg, "hi")
                if res:
                    from dotenv import set_key
                    set_key(".env", "AZURE_MAIN_KEY", key)
                    set_key(".env", "AZURE_ENDPOINT", ep)
                    set_key(".env", "AZURE_REGION", reg)
                    
                    os.environ["AZURE_MAIN_KEY"] = key
                    os.environ["AZURE_ENDPOINT"] = ep
                    os.environ["AZURE_REGION"] = reg
                    
                    self.after(0, lambda: self.log("✅ Azure OK!"))
                    self.after(0, lambda: messagebox.showinfo("OK", "Azure Key hợp lệ!"))
            except Exception as e:
                self.after(0, lambda: self.log(f"❌ Azure lỗi: {e}"))
        Thread(target=task, daemon=True).start()

    def handle_export(self):
        if not any(self.translated_lines):
            messagebox.showwarning("Lỗi", "Không có nội dung để xuất.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx"), ("Text", "*.txt")])
        if not path: 
            return
        try:
            if path.endswith(".docx"):
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                doc = Document()
                doc.add_heading('BẢN DỊCH AI MASTER', 0)
                for line in self.translated_lines:
                    if not line: 
                        continue
                    if utils.is_chapter_title(line):
                        p = doc.add_heading(line, level=1)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        doc.add_paragraph(line)
                doc.save(path)
            else:
                with open(path, "w", encoding="utf-8-sig") as f:
                    f.write("\n\n".join([l for l in self.translated_lines if l]))
            messagebox.showinfo("Xong", "Đã xuất file thành công!")
            self.log(f"💾 Đã xuất file: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))
            self.log(f"❌ Lỗi xuất file: {e}")

    def open_api(self):
        if not self.api_window or not self.api_window.winfo_exists():
            self.api_window = ApiManagerWindow(self)
        self.api_window.focus()
        self.log("🔧 Mở cửa sổ cấu hình API.")

    def stop_task(self):
        if not self.is_running:
            self.log("ℹ️ Không có tiến trình nào đang chạy.")
            return
        self.is_running = False
        self.log("🛑 Yêu cầu dừng đa luồng đã gửi. Hệ thống đang dọn dẹp và đóng Pool...")

    def refresh_ollama(self):
        def task():
            m = OllamaProvider.get_models()
            if m:
                self.after(0, lambda: self.model_selector.configure(values=m))
                self.after(0, lambda: self.model_selector.set(m[0]))
                self.after(0, lambda: self.log("🔄 Đã cập nhật danh sách model Ollama."))
            else:
                self.after(0, lambda: self.log("⚠️ Không tìm thấy model Ollama local hoặc chưa khởi động dịch vụ."))
        Thread(target=task, daemon=True).start()

if __name__ == "__main__":
    app = TranslatorApp()
    app.mainloop()
       
